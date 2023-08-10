"""
Loader utils.
"""

# Python built-in
from datetime import datetime
from io import BytesIO

import pytz
from docx import Document
import fitz


def decode_docx(content: bytes) -> str:
    stream = BytesIO(content)
    # stream = BytesIO(content)
    docx = Document(stream)
    full_text = []
    full_text = [para.text for para in docx.paragraphs]

    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    # text = docx.read('word/document.xml').decode('utf-8')
    # cleaned = re.sub(r'<(.|\n)*?>','',text)
    text = "\n".join(full_text).replace('\n\n','')
    return text[:text.find("Condiciones de privacidad")]


def decode_pdf(content: bytes) -> str:
    stream = BytesIO(content)
    doc = fitz.open(stream=stream, filetype="pdf")
    text= "".join(map(lambda page: page.get_textpage().extractText(),doc.pages(step=1)))
    # for page in doc.pages(step=1):
    #     text += page.get_textpage().extractText()
    cleaned = text.strip() #re.sub('\n\n','',text)
    return cleaned


def local_now() -> datetime:
    return datetime.now(pytz.timezone("UTC"))
